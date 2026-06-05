from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel
from pathlib import Path
from typing import List, Optional
import json
from app.api.projects import get_loader
from app.core.executor import Executor

router = APIRouter()

def _save_to_history(project_path, template, params, result):
    """Сохраняет результат анализа в историю проекта"""
    from datetime import datetime
    
    history_path = project_path / "state" / "analysis_history.json"
    
    # Загружаем существующую историю
    if history_path.exists():
        with open(history_path, 'r', encoding='utf-8') as f:
            history = json.load(f)
    else:
        history = []
    
    # Генерируем title
    title_map = {
        'cox_ph': f"Cox regression ({params.get('regression_type', 'multi')})",
        'logistic': f"Logistic regression ({params.get('regression_type', 'multi')})",
        'kaplan_meier': f"Kaplan-Meier ({params.get('plot_type', 'survival')})",
        'categorical': f"Categorical: {params.get('col1', '?')} vs {params.get('col2', '?')}",
        'random_forest': f"Random Forest (target: {params.get('target_col', '?')})",
        'roc_analysis': f"ROC Analysis ({params.get('target_col', '?')})",
        'model_evaluation_binary': f"Model Evaluation: {params.get('target_col', '?')}",
        'ai_chat': f"AI Chat ({params.get('title', 'Chat')})",
    }
    title = title_map.get(template, template)
    
    # Создаём запись
    record = {
        "id": f"run_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
        "timestamp": datetime.now().isoformat(),
        "template": template,
        "title": title,
        "params": {k: v for k, v in params.items() if not isinstance(v, (list, dict)) or len(str(v)) < 500},
        "metrics": result.get("metrics", {}),
        "plots": result.get("metrics", {}).get("plots", []) if result.get("metrics") else [],
        "status": "success" if result.get("success") else "error",
        "error": result.get("error") if not result.get("success") else None,
        "output_preview": result.get("output", "")[:2000]  # первые 2000 символов
    }
    
    # Добавляем в начало списка
    history.insert(0, record)
    
    # Сохраняем (без ограничения)
    with open(history_path, 'w', encoding='utf-8') as f:
        json.dump(history, f, indent=2, ensure_ascii=False, default=str)
    
    print(f"[HISTORY] Saved run {record['id']} to {history_path}")

class CodeRequest(BaseModel):
    code: str

class AnalysisRequest(BaseModel):
    template: str
    params: dict

class PredictRequest(BaseModel):
    template: str
    params: dict
    column_name: str = "predicted_risk"

@router.post("/code/run")
async def run_code(req: CodeRequest):
    loader = get_loader()
    if loader.df is None:
        raise HTTPException(status_code=400, detail="No data loaded")
    executor = Executor(loader.project_path)
    result = executor.execute_code(req.code)
    return result

@router.post("/run")
async def run_analysis(req: AnalysisRequest):
    loader = get_loader()
    if loader.df is None:
        raise HTTPException(status_code=400, detail="No data loaded")
    executor = Executor(loader.project_path)
    from app.core.modeling_orchestrator import ModelingOrchestrator
    orchestrator = ModelingOrchestrator(loader.project_path, executor)
    try:
        result = orchestrator.execute_template(req.template, req.params)
    except Exception as e:
        import traceback
        return {"success": False, "error": str(e), "traceback": traceback.format_exc(), "output": "", "metrics": {}, "table": []}
    
    _save_to_history(loader.project_path, req.template, req.params, result)
    
    return result

@router.post("/model/upload")
async def upload_model(file: UploadFile = File(...)):
    """Загружает файл модели и возвращает метаданные"""
    import joblib, tempfile, os
    
    loader = get_loader()
    if not loader.project_path:
        raise HTTPException(status_code=400, detail="No active project")
    
    # Save uploaded file
    models_dir = loader.project_path / "models"
    models_dir.mkdir(parents=True, exist_ok=True)
    file_path = models_dir / file.filename
    
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    # Load and extract info
    try:
        bundle = joblib.load(str(file_path))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid model file: {e}")
    
    if isinstance(bundle, dict):
        info = {
            "type": "sklearn_bundle",
            "features": bundle.get('features', []),
            "target": bundle.get('target', 'Unknown'),
            "encodings": bundle.get('feature_encodings', {}),
            "training_metrics": bundle.get('training_metrics', {}),
            "n_training": bundle.get('n_training', 0),
            "model_path": str(file_path)
        }
    else:
        info = {
            "type": type(bundle).__name__,
            "features": [],
            "target": "Unknown",
            "encodings": {},
            "training_metrics": {},
            "model_path": str(file_path)
        }
    
    return {"success": True, "info": info}

@router.post("/predict")
async def save_predictions(req: PredictRequest):
    """Сохраняет предсказанные значения как новую колонку в датафрейме"""
    loader = get_loader()
    if loader.df is None:
        raise HTTPException(status_code=400, detail="No data loaded")
    
    executor = Executor(loader.project_path)
    
    template = req.template
    params = req.params
    column_name = req.column_name
    
    if column_name in loader.df.columns:
        raise HTTPException(status_code=400, detail=f"Column '{column_name}' already exists")
    
    # Разная логика для Cox и Logistic
    if template in ["cox_ph", "cox"]:
        time_col = params.get('time_col', '')
        event_col = params.get('event_col', '')
        covariates = params.get('covariates', [])
        regression_type = params.get('regression_type', 'multi')
        reference_groups = params.get('reference_groups', {})
        covariate_types = params.get('covariate_types', {})
        p_enter = params.get('p_enter', 0.05)
        p_remove = params.get('p_remove', 0.10)
        
        if not time_col or not event_col:
            raise HTTPException(status_code=400, detail="Time and event columns required")
        
        code = f'''
import pandas as pd
import numpy as np
from lifelines import CoxPHFitter
from scipy.stats import chi2

time_col = "{time_col}"
event_col = "{event_col}"
covariates = {covariates}
column_name = "{column_name}"
regression_type = "{regression_type}"
reference_groups = {reference_groups}
covariate_types = {covariate_types}
p_enter = {p_enter}
p_remove = {p_remove}

print(f"Saving predicted risk to column: {{column_name}}")

df_pred = df.copy()

# Подготовка данных с dummy-кодированием
df_work = df_pred[[time_col, event_col] + covariates].dropna()

if len(df_work) < 10:
    raise ValueError("Insufficient data to train model")

encoding_map = {{}}
dummy_to_info = {{}}
encoded_covs = []

for col in covariates:
    if col in covariate_types:
        is_cat = (covariate_types[col] == 'categorical')
    else:
        is_cat = (df_work[col].dtype == 'object' or 
                  df_work[col].dtype.name == 'category' or 
                  isinstance(df_work[col].dtype, pd.CategoricalDtype))
    
    if is_cat:
        df_work[col] = df_work[col].astype(str).fillna('__MISSING__')
        categories = sorted(list(df_work[col].unique()))
        
        ref = reference_groups.get(col)
        if ref not in categories:
            ref = categories[0] if categories else '__MISSING__'
            
        encoding_map[col] = {{
            'categories': categories,
            'reference': ref,
            'dummy_cols': []
        }}
        
        for cat in categories:
            if cat == ref:
                continue
            dummy_name = f"{{col}}_{{cat}}"
            df_work[dummy_name] = (df_work[col] == cat).astype(int)
            # Также добавляем их в df_pred для этапа предсказания
            df_pred[dummy_name] = (df_pred[col].astype(str).fillna('__MISSING__') == cat).astype(int)
            encoding_map[col]['dummy_cols'].append(dummy_name)
            encoded_covs.append(dummy_name)
    else:
        # Cast to numeric to match training logic
        df_work[col] = pd.to_numeric(df_work[col], errors='coerce')
        df_pred[col] = pd.to_numeric(df_pred[col], errors='coerce')
        encoded_covs.append(col)

def fit_cox(covs_list):
    if not covs_list:
        return None
    cph = CoxPHFitter()
    cph.fit(df_work[[time_col, event_col] + covs_list], duration_col=time_col, event_col=event_col)
    return cph

selected_vars = covariates.copy()

# Пошаговый отбор, если требуется
if regression_type in ["forward", "backward"] and len(covariates) > 1:
    def get_model_info(vars_list):
        if not vars_list:
            return None
        covs = []
        for v in vars_list:
            if v in encoding_map:
                covs.extend(encoding_map[v]['dummy_cols'])
            else:
                covs.append(v)
        return fit_cox(covs)

    # Вычисляем null model log-likelihood
    first_var = covariates[0]
    first_cph = get_model_info([first_var])
    if first_cph:
        null_ll = first_cph.log_likelihood_ - 0.5 * first_cph.log_likelihood_ratio_test().test_statistic
    else:
        null_ll = -9999.0
        
    if regression_type == "forward":
        selected_vars = []
        remaining_vars = covariates.copy()
        
        while remaining_vars:
            best_p = 1.0
            best_var = None
            
            if not selected_vars:
                base_ll = null_ll
                base_df = 0
            else:
                base_model = get_model_info(selected_vars)
                base_ll = base_model.log_likelihood_ if base_model else null_ll
                base_df = len(base_model.params_) if base_model else 0
                
            for var in remaining_vars:
                test_model = get_model_info(selected_vars + [var])
                if test_model is None:
                    continue
                lrt_stat = 2 * (test_model.log_likelihood_ - base_ll)
                df_diff = len(test_model.params_) - base_df
                if df_diff > 0:
                    p_val = chi2.sf(lrt_stat, df=df_diff)
                    if p_val < best_p:
                        best_p = p_val
                        best_var = var
            
            if best_p < p_enter and best_var:
                selected_vars.append(best_var)
                remaining_vars.remove(best_var)
            else:
                break
    else: # backward
        selected_vars = covariates.copy()
        while len(selected_vars) > 1:
            worst_p = -1.0
            worst_var = None
            
            full_model = get_model_info(selected_vars)
            if full_model is None:
                break
            full_ll = full_model.log_likelihood_
            full_df = len(full_model.params_)
            
            for var in selected_vars:
                reduced_vars = [v for v in selected_vars if v != var]
                reduced_model = get_model_info(reduced_vars)
                reduced_ll = reduced_model.log_likelihood_ if reduced_model else null_ll
                reduced_df = len(reduced_model.params_) if reduced_model else 0
                
                lrt_stat = 2 * (full_ll - reduced_ll)
                df_diff = full_df - reduced_df
                if df_diff > 0:
                    p_val = chi2.sf(lrt_stat, df=df_diff)
                    if p_val > worst_p:
                        worst_p = p_val
                        worst_var = var
            
            if worst_p > p_remove and worst_var:
                selected_vars.remove(worst_var)
            else:
                break

# Формируем окончательный набор ковариат
final_covs = []
for v in selected_vars:
    if v in encoding_map:
        final_covs.extend(encoding_map[v]['dummy_cols'])
    else:
        final_covs.append(v)

cph = fit_cox(final_covs)

df_pred[column_name] = np.nan
valid_idx = df_pred[covariates].notna().all(axis=1)

if valid_idx.sum() > 0 and cph is not None:
    df_valid = df_pred.loc[valid_idx, final_covs].copy()
    partial_hazard = cph.predict_partial_hazard(df_valid)
    df_pred.loc[valid_idx, column_name] = partial_hazard.values

risk_col = df_pred[column_name]
if risk_col.notna().sum() > 0:
    min_risk = risk_col.min()
    max_risk = risk_col.max()
    if max_risk > min_risk:
        df_pred[column_name + "_norm"] = (risk_col - min_risk) / (max_risk - min_risk)
        print(f"Created normalized risk column: {{column_name}}_norm")

df = df_pred
print(f"Successfully added column '{{column_name}}'")
print(f"Model C-index: {{cph.concordance_index_:.4f}}")
'''
    
    elif template == "categorical":
        col1 = params.get('col1', '')
        col2 = params.get('col2', '')
        
        if not col1:
            raise HTTPException(status_code=400, detail="First variable required")
        if not col2:
            raise HTTPException(status_code=400, detail="Second variable required")
        
        result = orchestrator.execute_template("categorical_test", {"col1": col1, "col2": col2})
        return {"results": result, "output": result.get("output", ""), "success": result.get("success", False)}
    
    elif template in ["logistic", "logistic_regression"]:
        target_col = params.get('target_col', '')
        predictors = params.get('predictors', [])
        covariate_types = params.get('covariate_types', {})
        reference_groups = params.get('reference_groups', {})
        
        if not target_col:
            raise HTTPException(status_code=400, detail="Target column required")
        if not predictors:
            raise HTTPException(status_code=400, detail="Predictors required")
        
        code = f'''
import pandas as pd
import numpy as np
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler

target_col = "{target_col}"
predictors = {predictors}
column_name = "{column_name}"
covariate_types = {covariate_types}
reference_groups = {reference_groups}

print(f"Saving predicted probabilities to column: {{column_name}}")

df_pred = df.copy()

# Кодируем целевую переменную если нужно
if df_pred[target_col].dtype == 'object':
    from sklearn.preprocessing import LabelEncoder
    le_target = LabelEncoder()
    df_pred[target_col] = le_target.fit_transform(df_pred[target_col].astype(str))
    print(f"Target encoded: {{le_target.classes_}}")

# Dummy-кодирование категориальных предикторов
encoding_map = {{}}
dummy_to_info = {{}}
encoded_covs = []
for col in predictors:
    if col in covariate_types:
        is_cat = (covariate_types[col] == 'categorical')
    else:
        is_cat = (df_pred[col].dtype == 'object' or
                  df_pred[col].dtype.name == 'category' or
                  isinstance(df_pred[col].dtype, pd.CategoricalDtype))
    if is_cat:
        df_pred[col] = df_pred[col].astype(str).fillna('__MISSING__')
        categories = sorted(list(df_pred[col].unique()))
        ref = reference_groups.get(col)
        if ref not in categories:
            ref = categories[0] if categories else '__MISSING__'
        encoding_map[col] = {{'categories': categories, 'reference': ref, 'dummy_cols': []}}
        for cat in categories:
            if cat == ref:
                continue
            dummy_name = f"{{col}}_{{cat}}"
            df_pred[dummy_name] = (df_pred[col] == cat).astype(int)
            encoding_map[col]['dummy_cols'].append(dummy_name)
            encoded_covs.append(dummy_name)
    else:
        df_pred[col] = pd.to_numeric(df_pred[col], errors='coerce')
        encoded_covs.append(col)

# Удаляем пропуски для обучения
cols_for_model = [target_col] + encoded_covs
df_model = df_pred[cols_for_model].dropna()

print(f"Training on {{len(df_model)}} complete cases")

X = df_model[encoded_covs].values
y = df_model[target_col].values

scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

model = LogisticRegression(max_iter=1000, random_state=42)
model.fit(X_scaled, y)

# Предсказываем для всех строк
df_pred[column_name] = np.nan
valid_idx = df_pred[encoded_covs].notna().all(axis=1)

if valid_idx.sum() > 0:
    X_all = df_pred.loc[valid_idx, encoded_covs].values
    X_all_scaled = scaler.transform(X_all)
    probabilities = model.predict_proba(X_all_scaled)[:, 1]
    df_pred.loc[valid_idx, column_name] = probabilities

# Удаляем вспомогательные dummy-колонки
for col in encoding_map:
    for dcol in encoding_map[col]['dummy_cols']:
        if dcol in df_pred.columns:
            df_pred.drop(columns=[dcol], inplace=True)

df = df_pred
print(f"Successfully added column '{{column_name}}' with probabilities")
print(f"Model AUC: {{model.score(X_scaled, y):.4f}}")
'''
    
    elif template == "random_forest":
        target_col = params.get('target_col', '')
        exclusions = params.get('exclusions', '')
        
        if not target_col:
            raise HTTPException(status_code=400, detail="Target column required")
        
        code = f'''
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import LabelEncoder
from sklearn.model_selection import train_test_split

target_col = "{target_col}"
column_name = "{column_name}"

print(f"Saving predicted probabilities to column: {{column_name}}")

df_pred = df.copy()

# Автоисключение
user_excl_raw = "{exclusions}"
if user_excl_raw and user_excl_raw.lower() not in ['none', '[]', '']:
    user_exclusions = [x.strip() for x in user_excl_raw.split(',') if x.strip()]
else:
    user_exclusions = []

auto_exclude = []
for col in df_pred.columns:
    col_lower = col.lower()
    if any(p in col_lower for p in ['_event', 'event_', '_time', 'time_', 'predicted_', '_prob', 'censored', 'followup', 'cox_', 'logistic_', 'rf_']):
        auto_exclude.append(col)
auto_exclude = list(set(auto_exclude))
all_exclude = set(auto_exclude + user_exclusions)

# Ковариаты - все числовые кроме исключений
valid_cov = [c for c in df_pred.columns if c != target_col and pd.api.types.is_numeric_dtype(df_pred[c]) and c not in all_exclude]

if len(valid_cov) == 0:
    raise ValueError("No valid predictors")

# Удаляем пропуски
mask = df_pred[target_col].notna() & df_pred[valid_cov].notna().all(axis=1)
X = df_pred.loc[mask, valid_cov].copy()
y = df_pred.loc[mask, target_col].copy()

# Кодируем целевую
if y.dtype == 'object' or y.dtype.name == 'category':
    le = LabelEncoder()
    y = le.fit_transform(y)

print(f"Training on {{len(X)}} complete cases")

# Обучение
rf = RandomForestClassifier(n_estimators=200, max_depth=10, min_samples_leaf=5, random_state=42, n_jobs=-1)
rf.fit(X, y)

# Предсказания
df_pred[column_name] = np.nan
proba = rf.predict_proba(X)[:, 1]
df_pred.loc[X.index, column_name] = proba

df = df_pred
print(f"Successfully added column '{{column_name}}' with probabilities")
'''
    
    else:
        raise HTTPException(status_code=400, detail=f"Unknown template: {template}")
    
    result = executor.execute_code(code)
    
    if not result["success"]:
        raise HTTPException(status_code=500, detail=result.get("error", "Prediction failed"))
    
    # Обновляем список колонок
    from app.api.projects import get_loader as get_fresh_loader
    fresh_loader = get_fresh_loader()
    
    return {
        "success": True,
        "output": result["output"],
        "column_name": column_name,
        "columns": fresh_loader.get_columns_info()
    }

@router.get("/charts")
async def list_charts():
    loader = get_loader()
    plots_dir = loader.project_path / "plots"
    charts = []
    if plots_dir.exists():
        for png in plots_dir.glob("*.png"):
            charts.append((png.name, png.stat().st_mtime))
        charts.sort(key=lambda x: x[1], reverse=True)
        charts = [c[0] for c in charts]
    return {"charts": charts}

@router.get("/history")
async def get_history(limit: int = 50):
    """Получить историю анализов проекта"""
    loader = get_loader()
    history_path = loader.project_path / "state" / "analysis_history.json"
    
    if history_path.exists():
        with open(history_path, 'r', encoding='utf-8') as f:
            history = json.load(f)
        return {
            "history": history[:limit],
            "total": len(history)
        }
    return {"history": [], "total": 0}


@router.get("/history/{run_id}")
async def get_history_item(run_id: str):
    """Получить конкретную запись истории"""
    loader = get_loader()
    history_path = loader.project_path / "state" / "analysis_history.json"
    
    if history_path.exists():
        with open(history_path, 'r', encoding='utf-8') as f:
            history = json.load(f)
        
        for record in history:
            if record.get("id") == run_id:
                return record
    
    raise HTTPException(status_code=404, detail="Record not found")


@router.delete("/history")
async def clear_history():
    """Очистить историю проекта"""
    loader = get_loader()
    history_path = loader.project_path / "state" / "analysis_history.json"
    
    if history_path.exists():
        history_path.unlink()
        return {"status": "cleared"}
    return {"status": "already_empty"}


@router.delete("/plots")
async def clear_plots():
    """Очистить все графики проекта"""
    loader = get_loader()
    plots_dir = loader.project_path / "plots"
    
    count = 0
    if plots_dir.exists():
        for f in plots_dir.glob("*.png"):
            f.unlink()
            count += 1
    
    return {"status": "cleared", "count": count}


@router.post("/history/chat")
async def save_chat(req: dict):
    """Сохранить AI-диалог в историю"""
    loader = get_loader()
    
    _save_to_history(
        loader.project_path,
        'ai_chat',
        {
            'title': req.get('title', 'AI Chat'),
            'model': req.get('model', 'unknown'),
            'provider': req.get('provider', 'unknown'),
            'messages': req.get('messages', []),
            'n_analyses': req.get('n_analyses', 0),
            'n_columns': req.get('n_columns', 0)
        },
        {'success': True, 'metrics': {}, 'output': ''}
    )
    
    return {"status": "saved"}


class DocxReportRequest(BaseModel):
    title: str = "PDD_STAT Analysis Report"
    analyses: str = "all"
    selected_ids: List[str] = []
    include_sections: List[str] = ["overview", "tables", "metrics"]
    include_ai: bool = False

@router.post("/report/generate")
async def generate_docx_report(req: DocxReportRequest):
    """Генерирует DOCX отчёт на основе выбранных анализов"""
    from datetime import datetime
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import numpy as np
    
    loader = get_loader()
    history_path = loader.project_path / "state" / "analysis_history.json"
    
    if not history_path.exists():
        raise HTTPException(status_code=400, detail="No analysis history found")
    
    with open(history_path, 'r', encoding='utf-8') as f:
        all_history = json.load(f)
    
    analysis_history = [h for h in all_history if h.get('template') != 'ai_chat']
    
    if req.analyses == "all":
        selected = analysis_history
    else:
        selected = [h for h in analysis_history if h.get('id') in req.selected_ids]
    
    if not selected:
        raise HTTPException(status_code=400, detail="No analyses selected")
    
    doc = Document()
    
    # Настройка стилей документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0

    # Настройка заголовков
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Times New Roman'
        heading_font.color.rgb = RGBColor(0, 0, 0)
        heading_font.bold = True
        heading_style.paragraph_format.space_before = Pt(6)
        heading_style.paragraph_format.space_after = Pt(3)
    
    title = doc.add_heading(req.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Project: {loader.project_path.name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Analyses included: {len(selected)}")
    doc.add_paragraph()
    
    if "overview" in req.include_sections:
        doc.add_heading("Dataset Overview", level=1)
        df = loader.df
        if df is not None:
            cols = loader.get_columns_info()
            p = doc.add_paragraph()
            p.add_run(f"Rows: {len(df)}").bold = True
            p.add_run(f"  |  Columns: {len(cols)}")
            
            nums = [c['name'] for c in cols if c['type'] == 'numeric']
            bins = [c['name'] for c in cols if c['type'] == 'binary']
            cats = [c['name'] for c in cols if c['type'] == 'categorical']
            
            if nums:
                doc.add_paragraph(f"Numeric: {', '.join(nums[:10])}")
            if bins:
                doc.add_paragraph(f"Binary: {', '.join(bins[:5])}")
            if cats:
                doc.add_paragraph(f"Categorical: {', '.join(cats[:5])}")
        doc.add_paragraph()
    
    for idx, item in enumerate(selected, 1):
        template = item.get('template', 'unknown')
        title_text = item.get('title', template)
        timestamp = item.get('timestamp', '')[:16].replace('T', ' ')
        metrics = item.get('metrics', {})
        
        doc.add_heading(f"{idx}. {title_text}", level=1)
        doc.add_paragraph(f"Date: {timestamp}  |  Type: {template}")
        
        if "metrics" in req.include_sections:
            doc.add_heading("Key Metrics", level=2)
            metric_items = []
            for k, v in metrics.items():
                if k in ('model_type', 'plots', 'summary', 'number_at_risk', 
                        'survival_probability', 'selected_features', 'zeroed_features',
                        'top_features', 'results', 'auto_select_C', 'target', 'warnings'):
                    continue
                if isinstance(v, (int, float)) and v is not None:
                    if isinstance(v, float) and (np.isnan(v) or np.isinf(v)):
                        continue
                    v_str = f"{v:.4f}" if isinstance(v, float) else str(v)
                    metric_items.append(f"{k}: {v_str}")
                elif isinstance(v, str) and len(v) < 50:
                    metric_items.append(f"{k}: {v}")
            
            if metric_items:
                for m in metric_items[:15]:
                    doc.add_paragraph(m, style='List Bullet')
            
            sel = metrics.get('selected_features', [])
            if sel:
                doc.add_paragraph("Selected Features:")
                for f in sel[:10]:
                    doc.add_paragraph(f"{f['feature']} (OR={f['or']:.2f})", style='List Bullet')
            
            zer = metrics.get('zeroed_features', [])
            if zer:
                doc.add_paragraph(f"Removed: {', '.join(zer)}")
        
# Results tables
        if "tables" in req.include_sections:
            doc.add_heading("Results", level=2)
            output_text = item.get('output_preview', '')
            if output_text:
                lines = output_text.split('\n')
                table_data = []
                in_table = False
                
                for line in lines:
                    stripped = line.strip()
                    
                    # Начало таблицы
                    if stripped.startswith('|') and '---' not in stripped:
                        in_table = True
                        cells = [c.strip() for c in stripped.split('|') if c.strip()]
                        if cells and len(cells) >= 2:
                            table_data.append(cells)
                    elif in_table and not stripped.startswith('|'):
                        # Конец таблицы
                        if table_data and len(table_data) >= 1:
                            ncols = max(len(r) for r in table_data)
                            table = doc.add_table(rows=len(table_data), cols=ncols)
                            table.style = 'Table Grid'
                            # Настройка шрифта таблицы
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.paragraph_format.space_before = Pt(0)
                                        paragraph.paragraph_format.space_after = Pt(0)
                                        paragraph.paragraph_format.line_spacing = 1.0
                                        for run in paragraph.runs:
                                            run.font.name = 'Times New Roman'
                                            run.font.size = Pt(12)
                                            run.font.color.rgb = RGBColor(0, 0, 0)
                            for i, row in enumerate(table_data):
                                for j, cell_text in enumerate(row):
                                    cell = table.rows[i].cells[j]
                                    cell.text = cell_text
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(12)
                            doc.add_paragraph()
                        table_data = []
                        in_table = False
                
                # Последняя таблица
                if table_data and len(table_data) >= 1:
                    ncols = max(len(r) for r in table_data)
                    table = doc.add_table(rows=len(table_data), cols=ncols)
                    table.style = 'Table Grid'
                    # Настройка шрифта таблицы
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_before = Pt(0)
                                paragraph.paragraph_format.space_after = Pt(0)
                                paragraph.paragraph_format.line_spacing = 1.0
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            cell = table.rows[i].cells[j]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    reports_dir = loader.project_path / "state" / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    safe_name = req.title.replace(' ', '_').replace('/', '_')[:50]
    filename = f"{safe_name}_{timestamp}.docx"
    filepath = reports_dir / filename
    
    doc.save(str(filepath))
    
    return {
        "status": "generated",
        "filename": filename,
        "path": str(filepath),
        "analyses_count": len(selected)
    }


class AIReportRequest(BaseModel):
    title: str = "AI Analysis Report"
    analyses: str = "all"
    selected_ids: List[str] = []
    language: str = "English"


@router.post("/report/ai-generate")
async def generate_ai_report(req: AIReportRequest):
    """Генерирует AI-отчёт в DOCX на основе выбранных анализов"""
    from datetime import datetime
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    loader = get_loader()
    
    # 1. Check AI config
    from app.core.ai_clients import AIClientFactory
    from app.api.ai import load_config as load_ai_config, load_context
    
    ai_config = load_ai_config(loader)
    provider = ai_config.get('provider', '')
    if not provider:
        raise HTTPException(status_code=400, detail="AI not configured. Please configure AI in Chat Settings.")
    if provider == 'groq':
        groq_cfg = ai_config.get('groq', {})
        if not groq_cfg.get('api_key'):
            raise HTTPException(status_code=400, detail="AI not configured. Groq API key missing.")
    if provider == 'ollama':
        ollama_cfg = ai_config.get('ollama', {})
        if not ollama_cfg.get('url'):
            raise HTTPException(status_code=400, detail="AI not configured. Ollama URL missing.")
    
    model = ai_config.get('last_used_model', 'llama3:8b')
    temperature = ai_config.get('temperature', 0.3)
    
    # 2. Load history
    history_path = loader.project_path / "state" / "analysis_history.json"
    if not history_path.exists():
        raise HTTPException(status_code=400, detail="No analysis history found")
    
    with open(history_path, 'r', encoding='utf-8') as f:
        all_history = json.load(f)
    
    analysis_history = [h for h in all_history if h.get('template') != 'ai_chat']
    
    if req.analyses == "all":
        selected = analysis_history
    else:
        selected = [h for h in analysis_history if h.get('id') in req.selected_ids]
    
    if not selected:
        raise HTTPException(status_code=400, detail="No analyses selected")
    
    # Cap at 5 analyses to keep prompt within token limits
    if len(selected) > 5:
        selected = selected[:5]
    
    # 3. Build system prompt for AI report generation
    import math
    
    def fmt_item(item):
        title = item.get('title', item.get('template', ''))
        ts = item.get('timestamp', '')[:10]
        m = item.get('metrics', {})
        parts = []
        for k, v in m.items():
            if k in ('model_type', 'plots', 'summary', 'auto_select_C', 'target', 'table', 'coefficients_list'):
                continue
            if isinstance(v, (int, float)) and v is not None:
                if isinstance(v, float) and (math.isnan(v) or math.isinf(v)):
                    continue
                vs = str(int(v)) if v == int(v) else f"{v:.4f}"
                parts.append(f"{k}={vs}")
            elif k == 'coefficients' and isinstance(v, list):
                rl = 'HR' if 'cox' in str(m.get('model_type','')).lower() else 'OR'
                cv = [f"{c.get('variable','?')}({rl}={c.get('hr') or c.get('or',0):.2f},p={c.get('p_value',1):.4f})" for c in v[:5]]
                if cv: parts.append(f"coeff: {'; '.join(cv)}")
            elif k == 'results' and isinstance(v, list):
                for r in v[:3]:
                    parts.append(f"{r.get('test') or r.get('predictor','?')}(Se={r.get('sensitivity',0):.2f},Sp={r.get('specificity',0):.2f})")
            elif k == 'schoenfeld_test' and isinstance(v, dict):
                sv = [f"{vn}(PH_p={rv.get('p_value',1):.4f})" for vn, rv in v.items()]
                if sv: parts.append(f"Schoenfeld: {'; '.join(sv[:3])}")
            elif k == 'vif' and isinstance(v, dict):
                vp = [f"{vn}(VIF={vi.get('value',0):.1f})" for vn, vi in v.items()]
                if vp: parts.append(f"VIF: {'; '.join(vp[:5])}")
            elif k == 'warnings' and isinstance(v, list) and v:
                parts.append(f"warn: {'; '.join(str(w)[:60] for w in v[:2])}")
            elif k == 'feature_importance' and isinstance(v, dict):
                sf = sorted(v.items(), key=lambda x: x[1], reverse=True)[:3]
                parts.append(f"imp: {'; '.join(f'{n}({x:.4f})' for n,x in sf)}")
            elif k == 'median_survival' and isinstance(v, dict):
                parts.append(f"medians: {'; '.join(f'{g}={val:.1f}mo' if val else f'{g}=NR' for g,val in v.items())}")
            elif k == 'metrics_by_model' and isinstance(v, dict):
                parts.append('; '.join(f"{mn}(AUC={mm.get('auc',0):.3f})" for mn, mm in v.items()))
            elif k == 'model_comparisons' and isinstance(v, list):
                parts.append('; '.join(f"{c['model1']}vs{c['model2']}(ΔC={c.get('delta',0):+.4f},p={c.get('p_value',1):.4f})" for c in v[:2]))
        return f"[{ts}] {title}: {'; '.join(parts[:12])}"
    
    ctx = load_context(loader)
    
    prompt_lines = [
        "You are a biostatistics report writer. Write a structured scientific report.",
        f"Language: {req.language}.",
        "",
        "STRUCTURE (use ## and ### markdown):",
        "1. Objective — from project context",
        "2. Materials — N cases, dataset overview",
        "3. Methods — statistical methods used",
        "4. Results — key numbers (HR, OR, p, AUC, C-index)",
        "5. Conclusion",
        "",
    ]
    if ctx.get('description'):
        prompt_lines.append(f"Project: {ctx['description'][:300]}")
    if ctx.get('aim'):
        prompt_lines.append(f"Aim: {ctx['aim'][:200]}")
    
    df = loader.df
    if df is not None:
        cols = loader.get_columns_info()
        nums = [c['name'] for c in cols if c['type'] == 'numeric'][:10]
        cats = [c['name'] for c in cols if c['type'] == 'categorical'][:3]
        bins = [c['name'] for c in cols if c['type'] == 'binary'][:3]
        prompt_lines.append(f"\nData: {len(df)} rows, {len(cols)} cols")
        if nums: prompt_lines.append(f"Numeric: {', '.join(nums)}")
        if bins: prompt_lines.append(f"Binary: {', '.join(bins)}")
        if cats: prompt_lines.append(f"Categorical: {', '.join(cats)}")
    
    prompt_lines.append("\n## Analyses:")
    for item in selected:
        prompt_lines.append(fmt_item(item))
    
    prompt_lines.append("\nIMPORTANT: Do NOT include <think> blocks or any internal reasoning. Output ONLY the report content directly.")
    prompt_lines.append("\nWrite the full report now with exact numbers.")
    prompt = '\n'.join(prompt_lines)
    
    # 5. Call AI
    max_tokens = 4000
    # Estimate tokens: ~4 chars per token
    if len(prompt) > 25000:
        raise HTTPException(status_code=400, 
            detail=f"Prompt too large ({len(prompt)} chars). Reduce number of selected analyses.")
    
    try:
        client = AIClientFactory.create(ai_config)
        result = await client.chat(
            model=model,
            messages=[{"role": "system", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens
        )
    except Exception as e:
        err = str(e)
        if 'rate_limit_exceeded' in err or 'Request too large' in err:
            raise HTTPException(status_code=400,
                detail=f"AI request exceeds token limit for {provider}. Try reducing number of analyses or switch to a model with larger context.")
        raise HTTPException(status_code=500, detail=f"AI request failed: {err}")
    
    if not result.get('success'):
        raise HTTPException(status_code=500, detail=result.get('error', 'AI request failed'))
    
    ai_content = result.get('content', '')
    if not ai_content.strip():
        raise HTTPException(status_code=500, detail="AI returned empty response")
    
    # Strip any <think> blocks (internal reasoning)
    ai_content = re.sub(r'<think>.*?</think>', '', ai_content, flags=re.DOTALL).strip()
    # Strip lone <think> without closing tag
    ai_content = re.sub(r'<think>.*', '', ai_content).strip()
    
    # 6. Convert AI markdown response to DOCX
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0
    
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(6)
        heading_style.paragraph_format.space_after = Pt(3)
    
    title_h = doc.add_heading(req.title, 0)
    title_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Project: {loader.project_path.name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Language: {req.language}")
    doc.add_paragraph(f"Analyses included: {len(selected)}")
    doc.add_paragraph()
    
    def add_markdown_paragraph(doc_obj, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        p = doc_obj.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            elif part.strip():
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        return p
    
    lines = ai_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('|') and line.endswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                if cells:
                    table_rows.append(cells)
                i += 1
            if table_rows:
                if len(table_rows) > 1 and all(c in '-: ' for c in table_rows[1][0].replace('|', '').replace('-', '').replace(':', '').strip()):
                    table_rows.pop(1)
                ncols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(table_rows):
                    for cj, cell_text in enumerate(row_data):
                        if cj < ncols and ri < len(table.rows):
                            cell = table.rows[ri].cells[cj]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                doc.add_paragraph()
            continue
        elif line.strip() == '':
            doc.add_paragraph('')
        elif line.startswith('- ') and len(line) > 2:
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        else:
            add_markdown_paragraph(doc, line)
        
        i += 1
    
    # 7. Save
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    reports_dir = loader.project_path / "state" / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    safe_name = req.title.replace(' ', '_').replace('/', '_')[:50]
    filename = f"AI_{safe_name}_{timestamp}.docx"
    filepath = reports_dir / filename
    
    doc.save(str(filepath))
    
    return {
        "status": "generated",
        "filename": filename,
        "path": str(filepath),
        "analyses_count": len(selected)
    }

@router.get("/report/download/{filename}")
async def download_report(filename: str):
    """Скачать сгенерированный отчёт"""
    from fastapi.responses import FileResponse
    loader = get_loader()
    filepath = loader.project_path / "state" / "reports" / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)


class ReportRequest(BaseModel):
    template: str
    title: str
    output: str
    metrics: dict
    charts: list


@router.post("/report")
async def save_report(req: ReportRequest):
    from app.api import projects as proj_module
    pm = proj_module.get_project_manager()
    if not pm.current_project_path:
        raise HTTPException(status_code=400, detail="No active project")

    reports_dir = pm.current_project_path / "state" / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    report_name = f"{pm.current_project_name}_report_{req.template}.html"
    report_path = reports_dir / report_name

    def md_to_html(text):
        import re
        html = text
        html = html.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', html)
        html = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img alt="\1" src="\2" />', html)
        html = re.sub(r'`([^`]+)`', r'<code>\1</code>', html)
        return html

    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{req.title}</title>
    <style>
        * {{ box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; padding: 30px; max-width: 900px; margin: 0 auto; background: #fafafa; color: #333; line-height: 1.6; }}
        .header {{ background: #fff; padding: 24px; border-radius: 12px; margin-bottom: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        h1 {{ margin: 0 0 8px 0; font-size: 22px; color: #1a1a1a; }}
        .date {{ color: #888; font-size: 13px; }}
        .metrics {{ background: #fff; padding: 20px; border-radius: 12px; margin-bottom: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .metrics-title {{ font-weight: 600; font-size: 16px; margin-bottom: 12px; }}
        .metric {{ display: inline-block; margin-right: 24px; margin-bottom: 8px; }}
        .metric-label {{ color: #666; font-size: 12px; }}
        .metric-value {{ font-size: 24px; font-weight: 600; color: #2563eb; }}
        .results {{ background: #fff; padding: 20px; border-radius: 12px; margin-bottom: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        table {{ width: 100%; border-collapse: collapse; margin: 12px 0; }}
        th, td {{ padding: 10px 12px; text-align: left; border-bottom: 1px solid #eee; }}
        th {{ background: #f8f9fa; font-weight: 600; font-size: 13px; color: #555; }}
        td {{ font-size: 14px; }}
        tr:hover {{ background: #f8f9fa; }}
        .chart {{ background: #fff; padding: 20px; border-radius: 12px; text-align: center; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .chart img {{ max-width: 100%; border-radius: 8px; }}
        code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 4px; font-size: 13px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{req.title}</h1>
        <div class="date">Project: {pm.current_project_name}</div>
    </div>
    <div class="metrics">
        <div class="metrics-title">Model Metrics</div>
        {''.join(f'<div class="metric"><div class="metric-label">{k}</div><div class="metric-value">{v}</div></div>' for k, v in req.metrics.items())}
    </div>
    <div class="results">
        {md_to_html(req.output)}
    </div>
    {''.join(f'<div class="chart"><img src="../plots/{c}" /></div>' for c in req.charts) if req.charts else ''}
</body>
</html>"""

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html)

    return {"status": "saved", "path": str(report_path)}


@router.get("/reports")
async def list_reports():
    from app.api import projects as proj_module
    pm = proj_module.get_project_manager()
    if not pm.current_project_path:
        return {"reports": []}

    reports_dir = pm.current_project_path / "state" / "reports"
    reports = []
    if reports_dir.exists():
        for f in reports_dir.glob("*_report*.html"):
            reports.append({"name": f.name, "modified": f.stat().st_mtime})
    return {"reports": sorted(reports, key=lambda x: x['modified'], reverse=True)}


@router.get("/report/{name}")
async def get_report(name: str):
    from app.api import projects as proj_module
    from fastapi.responses import FileResponse
    pm = proj_module.get_project_manager()
    report_path = pm.current_project_path / "state" / "reports" / name
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    return FileResponse(report_path, media_type="text/html")
